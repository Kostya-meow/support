"""
Простые и быстрые парсеры документов
Поддерживает: PDF (OCR), DOCX, DOC, Markdown
"""

from __future__ import annotations

import logging
import os
import tempfile
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def resize_image_to_max_side(image, max_side: int = 1024):
    """Ресайз изображения до максимальной стороны"""
    from PIL import Image
    
    width, height = image.size
    
    # Если уже меньше - не ресайзим
    if max(width, height) <= max_side:
        return image
    
    # Вычисляем новые размеры с сохранением пропорций
    if width > height:
        new_width = max_side
        new_height = int(height * max_side / width)
    else:
        new_height = max_side
        new_width = int(width * max_side / height)
    
    return image.resize((new_width, new_height), Image.LANCZOS)


class DocumentParser:
    """Базовый класс для парсеров документов"""

    @staticmethod
    def extract_text(file_path: str | Path) -> str:
        """Извлечь текст из документа"""
        raise NotImplementedError


class PDFParser(DocumentParser):
    """Быстрый PDF парсер с OCR"""

    @staticmethod
    def extract_text(file_path: str | Path) -> str:
        """
        Простой и быстрый парсинг PDF:
        1. Каждую страницу конвертируем в картинку
        2. Ресайзим до 1024 макс стороны
        3. OCR через easyocr
        """
        try:
            from PIL import Image
            import easyocr
            import pdf2image
        except ImportError as e:
            logger.error(f"Missing dependencies: {e}")
            logger.error("Install with: pip install pillow easyocr pdf2image")
            raise ImportError("Required packages: pillow, easyocr, pdf2image")

        text_parts = []
        
        try:
            # Инициализируем OCR (только русский и английский для скорости)
            reader = easyocr.Reader(['ru', 'en'], gpu=False)
            
            # Конвертируем PDF в изображения
            logger.info(f"Converting PDF to images: {file_path}")
            images = pdf2image.convert_from_path(
                file_path,
                dpi=200,  # Повышаем качество для лучшего OCR
                first_page=None,
                last_page=None,
                poppler_path=None  # Используем системный poppler
            )
            
            logger.info(f"Processing {len(images)} pages with OCR")
            
            for page_num, image in enumerate(images, 1):
                try:
                    # Ресайзим для скорости
                    resized_image = resize_image_to_max_side(image, max_side=1024)
                    
                    # Конвертируем PIL Image в numpy array для easyocr
                    import numpy as np
                    image_array = np.array(resized_image)
                    
                    # OCR
                    results = reader.readtext(image_array, paragraph=True)
                    
                    # Собираем текст со страницы
                    page_text = []
                    for result in results:
                        # easyocr может возвращать разные форматы в зависимости от версии
                        if len(result) == 3:
                            # Формат: (bbox, text, confidence)
                            _, text, confidence = result
                        elif len(result) == 2:
                            # Формат: (bbox, text) - без confidence
                            _, text = result
                            confidence = 1.0  # Принимаем все результаты если нет confidence
                        else:
                            logger.warning(f"Unexpected OCR result format: {result}")
                            continue
                            
                        if confidence > 0.2:  # Снижаем порог для большего извлечения текста
                            clean_text = text.strip()
                            if clean_text and len(clean_text) > 1:  # Игнорируем слишком короткий текст
                                page_text.append(clean_text)
                    
                    if page_text:
                        page_content = " ".join(page_text)
                        text_parts.append(f"--- Страница {page_num} ---\n{page_content}")
                        logger.info(f"Page {page_num}: extracted {len(page_content)} chars")
                    else:
                        logger.warning(f"Page {page_num}: no text detected")
                    
                except Exception as e:
                    logger.warning(f"Error processing page {page_num}: {e}")
                    continue
                    
            final_text = "\n\n".join(text_parts)
            
            if not final_text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")
                return "PDF обработан, но текст не обнаружен. Возможно, файл содержит только изображения без текста."
                
            return final_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise


class DOCXParser(DocumentParser):
    """Простой DOCX парсер"""

    @staticmethod
    def extract_text(file_path: str | Path) -> str:
        """Простое извлечение текста из DOCX"""
        try:
            import docx
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ImportError("python-docx is required for DOCX parsing")

        try:
            doc = docx.Document(file_path)
            
            # Просто берем все параграфы
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # И все таблицы
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            
            all_text = "\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n--- Таблицы ---\n" + "\n".join(tables_text)
                
            return all_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise


class DOCParser(DocumentParser):
    """Простой DOC парсер (старые файлы Word)"""

    @staticmethod
    def extract_text(file_path: str | Path) -> str:
        """
        Извлечение текста из .doc файлов через docx2txt
        """
        try:
            import docx2txt
        except ImportError:
            logger.error("docx2txt not installed. Install with: pip install docx2txt")
            raise ImportError(
                "DOC files require docx2txt library. "
                "Install with: pip install docx2txt"
            )

        try:
            # docx2txt может работать с .doc файлами через antiword
            text = docx2txt.process(file_path)
            
            if not text or not text.strip():
                # Если не получилось через docx2txt, предлагаем конвертацию
                raise ValueError(
                    "Could not extract text from DOC file. "
                    "Please convert to DOCX format using LibreOffice or MS Word"
                )
                
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error parsing DOC {file_path}: {e}")
            # Fallback - предлагаем конвертацию
            raise ValueError(
                f"Failed to process DOC file: {e}. "
                "Please convert to DOCX format using LibreOffice or MS Word"
            )


class MarkdownParser(DocumentParser):
    """Простой Markdown парсер"""

    @staticmethod
    def extract_text(file_path: str | Path) -> str:
        """Простое чтение Markdown как есть"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return content
            
        except Exception as e:
            logger.error(f"Error parsing Markdown {file_path}: {e}")
            raise


class TXTParser(DocumentParser):
    """Простой TXT парсер"""

    @staticmethod
    def extract_text(file_path: str | Path) -> str:
        """Чтение обычных текстовых файлов"""
        try:
            # Пробуем несколько кодировок
            encodings = ['utf-8', 'windows-1251', 'cp1252', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    logger.info(f"Successfully read TXT file with {encoding} encoding")
                    return content
                except UnicodeDecodeError:
                    continue
                    
            # Если все кодировки не сработали, читаем как binary и пытаемся декодировать
            with open(file_path, "rb") as f:
                raw_content = f.read()
                content = raw_content.decode('utf-8', errors='ignore')
                logger.warning(f"TXT file read with error handling: {file_path}")
                return content
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {e}")
            raise


class DocumentParserFactory:
    """Простая фабрика парсеров"""

    _parsers = {
        ".pdf": PDFParser,
        ".docx": DOCXParser,
        ".doc": DOCParser,
        ".md": MarkdownParser,
        ".markdown": MarkdownParser,
        ".txt": TXTParser,
    }

    @classmethod
    def get_parser(cls, file_path: str | Path) -> DocumentParser:
        """Получить парсер для файла"""
        ext = Path(file_path).suffix.lower()
        
        parser_class = cls._parsers.get(ext)
        if not parser_class:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported: {', '.join(cls._parsers.keys())}"
            )
            
        return parser_class()

    @classmethod
    def parse_document(cls, file_path: str | Path) -> str:
        """Быстро парсит документ и возвращает текст"""
        logger.info(f"Parsing document: {file_path}")
        parser = cls.get_parser(file_path)
        text = parser.extract_text(file_path)
        logger.info(f"Extracted {len(text)} characters from {file_path}")
        return text

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """Список поддерживаемых расширений"""
        return list(cls._parsers.keys())
